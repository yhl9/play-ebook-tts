"""
DOCX文件处理器
"""

from typing import Optional, List
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from utils.log_manager import LogManager


class DOCXProcessor:
    """DOCX文件处理器"""
    
    def __init__(self):
        self.logger = LogManager().get_logger("DOCXProcessor")
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx未安装，DOCX处理功能不可用")
    
    def extract_text(self, file_path: str) -> str:
        """提取DOCX文本"""
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx未安装，无法处理DOCX文件")
            
            if not Path(file_path).exists():
                raise FileNotFoundError(f"DOCX文件不存在: {file_path}")
            
            self.logger.info(f"开始提取DOCX文本: {file_path}")
            
            # 打开DOCX文件
            doc = Document(file_path)
            
            text_content = ""
            
            # 提取段落文本，保持段落结构
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 检查段落是否为空行或只包含空白字符
                    if paragraph.text.strip():
                        text_content += paragraph.text.strip() + "\n\n"  # 使用双换行符分隔段落
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " ".join(row_text) + "\n\n"  # 表格行也用双换行符分隔
            
            if not text_content.strip():
                self.logger.warning("DOCX文件中未找到可提取的文本")
                return ""
            
            # 清理文本
            text_content = self._clean_extracted_text(text_content)
            
            self.logger.info(f"DOCX文本提取完成，文本长度: {len(text_content)}")
            return text_content
            
        except Exception as e:
            self.logger.error(f"DOCX文本提取失败: {file_path}, 错误: {e}")
            raise DOCXProcessingError(f"DOCX文本提取失败: {e}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """清理提取的文本"""
        if not text:
            return ""
        
        import re
        
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除多余的换行符
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 移除首尾空白
        text = text.strip()
        
        return text
    
    def get_docx_info(self, file_path: str) -> dict:
        """获取DOCX文件信息"""
        try:
            if not DOCX_AVAILABLE:
                return {}
            
            doc = Document(file_path)
            
            info = {
                'title': '',
                'author': '',
                'subject': '',
                'keywords': '',
                'created': '',
                'modified': '',
                'num_paragraphs': 0,
                'num_tables': 0,
                'num_pages': 0
            }
            
            # 获取文档属性
            core_props = doc.core_properties
            
            if core_props.title:
                info['title'] = core_props.title
            
            if core_props.author:
                info['author'] = core_props.author
            
            if core_props.subject:
                info['subject'] = core_props.subject
            
            if core_props.keywords:
                info['keywords'] = core_props.keywords
            
            if core_props.created:
                info['created'] = core_props.created.isoformat()
            
            if core_props.modified:
                info['modified'] = core_props.modified.isoformat()
            
            # 统计段落和表格数量
            info['num_paragraphs'] = len(doc.paragraphs)
            info['num_tables'] = len(doc.tables)
            
            # 估算页数（基于段落数量）
            info['num_pages'] = max(1, len(doc.paragraphs) // 20)
            
            return info
            
        except Exception as e:
            self.logger.error(f"获取DOCX信息失败: {file_path}, 错误: {e}")
            return {}
    
    def get_paragraph_list(self, file_path: str) -> List[dict]:
        """获取段落列表"""
        try:
            if not DOCX_AVAILABLE:
                return []
            
            doc = Document(file_path)
            paragraphs = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    paragraphs.append({
                        'index': i,
                        'text': paragraph.text.strip(),
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'length': len(paragraph.text)
                    })
            
            return paragraphs
            
        except Exception as e:
            self.logger.error(f"获取DOCX段落列表失败: {file_path}, 错误: {e}")
            return []
    
    def get_table_list(self, file_path: str) -> List[dict]:
        """获取表格列表"""
        try:
            if not DOCX_AVAILABLE:
                return []
            
            doc = Document(file_path)
            tables = []
            
            for i, table in enumerate(doc.tables):
                table_info = {
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'cells': len(table.rows) * len(table.columns)
                }
                
                # 提取表格内容
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        row_content.append(cell.text.strip())
                    table_content.append(row_content)
                
                table_info['content'] = table_content
                tables.append(table_info)
            
            return tables
            
        except Exception as e:
            self.logger.error(f"获取DOCX表格列表失败: {file_path}, 错误: {e}")
            return []
    
    def extract_paragraph_text(self, file_path: str, paragraph_index: int) -> str:
        """提取指定段落的文本"""
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx未安装，无法处理DOCX文件")
            
            doc = Document(file_path)
            
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引超出范围: {paragraph_index}")
            
            paragraph = doc.paragraphs[paragraph_index]
            return paragraph.text.strip()
            
        except Exception as e:
            self.logger.error(f"提取DOCX段落文本失败: {file_path}, 段落: {paragraph_index}, 错误: {e}")
            raise DOCXProcessingError(f"提取DOCX段落文本失败: {e}")
    
    def extract_table_text(self, file_path: str, table_index: int) -> str:
        """提取指定表格的文本"""
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx未安装，无法处理DOCX文件")
            
            doc = Document(file_path)
            
            if table_index < 0 or table_index >= len(doc.tables):
                raise ValueError(f"表格索引超出范围: {table_index}")
            
            table = doc.tables[table_index]
            table_text = ""
            
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text += " ".join(row_text) + "\n"
            
            return table_text.strip()
            
        except Exception as e:
            self.logger.error(f"提取DOCX表格文本失败: {file_path}, 表格: {table_index}, 错误: {e}")
            raise DOCXProcessingError(f"提取DOCX表格文本失败: {e}")
    
    def is_docx_valid(self, file_path: str) -> bool:
        """检查DOCX文件是否有效"""
        try:
            if not DOCX_AVAILABLE:
                return False
            
            doc = Document(file_path)
            return doc is not None
            
        except Exception as e:
            self.logger.error(f"检查DOCX文件有效性失败: {file_path}, 错误: {e}")
            return False
    
    def get_document_structure(self, file_path: str) -> dict:
        """获取文档结构"""
        try:
            if not DOCX_AVAILABLE:
                return {}
            
            doc = Document(file_path)
            
            structure = {
                'paragraphs': [],
                'tables': [],
                'headings': []
            }
            
            # 分析段落
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_info = {
                        'index': i,
                        'text': paragraph.text.strip(),
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'is_heading': paragraph.style.name.startswith('Heading') if paragraph.style else False
                    }
                    
                    structure['paragraphs'].append(para_info)
                    
                    # 如果是标题，添加到标题列表
                    if para_info['is_heading']:
                        structure['headings'].append(para_info)
            
            # 分析表格
            for i, table in enumerate(doc.tables):
                table_info = {
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns)
                }
                structure['tables'].append(table_info)
            
            return structure
            
        except Exception as e:
            self.logger.error(f"获取DOCX文档结构失败: {file_path}, 错误: {e}")
            return {}


class DOCXProcessingError(Exception):
    """DOCX处理异常"""
    pass
